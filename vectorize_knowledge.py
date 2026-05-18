import os
import glob
import openpyxl
from docx import Document
from pypdf import PdfReader
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

# DATA PATH CONFIGURATIONS
SOURCE_DIR = os.path.expanduser("~/talexcel_ai_bot/knowledge_source")
DB_DIR = os.path.expanduser("~/talexcel_ai_bot/chroma_db")

def parse_pdf(file_path):
    """Extract structural text from PDF files."""
    text_content = []
    try:
        reader = PdfReader(file_path)
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    except Exception as e:
        print(f"[ERROR] Failed to parse PDF {os.path.basename(file_path)}: {e}")
    return "\n".join(text_content)

def parse_docx(file_path):
    """Extract structural text from Microsoft Word documents."""
    text_content = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
    except Exception as e:
        print(f"[ERROR] Failed to parse DOCX {os.path.basename(file_path)}: {e}")
    return "\n".join(text_content)

def parse_xlsx(file_path):
    """Extract and linearize data from Microsoft Excel sheets into indexable text strings."""
    text_content = []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            text_content.append(f"=== Sheet: {sheet} ===")
            for row in ws.iter_rows(values_only=True):
                # Filter out completely empty rows, convert values to strings, and join with pipes
                row_values = [str(cell).strip() for cell in row if cell is not None]
                if row_values:
                    text_content.append(" | ".join(row_values))
    except Exception as e:
        print(f"[ERROR] Failed to parse XLSX {os.path.basename(file_path)}: {e}")
    return "\n".join(text_content)

def main():
    print("=== Starting TalExcel Knowledge Base Vectorization Pipeline ===")
    
    # 1. Initialize modern warning-free LangChain Embedding Architecture
    print("[INFO] Initializing HuggingFace Embeddings (bge-small-en-v1.5)...")
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5", model_kwargs={'device': 'cpu'})
    
    # 2. Set up the Recursive Character Splitter Strategy (750 chunk size / 100 overlap)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=750,
        chunk_overlap=100,
        separators=["\n\n", "\n", " ", ""]
    )
    
    # 3. Create ingestion storage folder if it doesn't exist
    if not os.path.exists(SOURCE_DIR):
        os.makedirs(SOURCE_DIR)
        print(f"[INIT] Created clean source folder at '{SOURCE_DIR}'. Place files there and re-run.")
        return

    # Scan and filter target file extensions
    extensions = ('*.pdf', '*.docx', '*.xlsx')
    discovered_files = []
    for ext in extensions:
        discovered_files.extend(glob.glob(os.path.join(SOURCE_DIR, ext)))
        
    if not discovered_files:
        print(f"[WARN] No matching files (.pdf, .docx, .xlsx) found inside target directory: {SOURCE_DIR}")
        return

    print(f"[INFO] Discovered {len(discovered_files)} documents for vector ingestion processing.")
    
    processed_documents = []
    
    # 4. Extract text based on file extensions
    for file_path in discovered_files:
        file_name = os.path.basename(file_path)
        print(f"[PARSING] Processing: {file_name}")
        
        extracted_text = ""
        if file_path.endswith('.pdf'):
            extracted_text = parse_pdf(file_path)
        elif file_path.endswith('.docx'):
            extracted_text = parse_docx(file_path)
        elif file_path.endswith('.xlsx'):
            extracted_text = parse_xlsx(file_path)
            
        if not extracted_text.strip():
            print(f"[SKIP] No readable text found or extracted from {file_name}")
            continue
            
        # Wrap the raw string inside a structural LangChain Document model along with tracking metadata
        doc_metadata = {"source": file_name, "path": file_path}
        processed_documents.append(LangChainDocument(page_content=extracted_text, metadata=doc_metadata))

    if not processed_documents:
        print("[ERROR] Database parsing yielded 0 valid document models. Aborting commit.")
        return

    # 5. Segment documents into optimized RAG text chunks
    print(f"[CHUNKING] Segmenting {len(processed_documents)} master documents into text nodes...")
    segmented_chunks = text_splitter.split_documents(processed_documents)
    print(f"[CHUNKING] Split complete. Total generated matrix sub-chunks: {len(segmented_chunks)}")

    # 6. Commit chunks directly to the persistent local ChromaDB folder
    print(f"[DATABASE] Writing coordinates and serialized mappings to vector cache at: {DB_DIR}")
    vector_db = Chroma.from_documents(
        documents=segmented_chunks,
        embedding=embeddings,
        persist_directory=DB_DIR
    )
    
    print("=== [SUCCESS] Ingestion Pipeline Finished. Phase 2 Vectorization Fully Built! ===")

if __name__ == "__main__":
    main()